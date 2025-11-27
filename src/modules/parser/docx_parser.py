from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Dict

from docx import Document  # библиотека python-docx

from .base import BaseParser
from .models import ParsedDocument
from .utils import normalize_text


class DocxParser(BaseParser):
    """
    Парсер документов Microsoft Word (.docx).
    """

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() == ".docx"

    def parse(self, path: Path) -> ParsedDocument:
        if not path.exists():
            raise FileNotFoundError(path)

        doc = Document(str(path))

        # тексты всех непустых параграфов
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)

        stat = path.stat()
        metadata: Dict[str, str] = {
            "file_name": path.name,
            "absolute_path": str(path.resolve()),
            "size_bytes": str(stat.st_size),
            "last_modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "content_type": (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        }

        return ParsedDocument(
            id=str(path.resolve()),
            raw_text=normalize_text(text),
            metadata=metadata,
        )
